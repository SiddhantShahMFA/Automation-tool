from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_FILE = "AI_Native_Delivery_Handbook.docx"


# ---------- Helpers ----------

def set_page_margins(section, top=0.7, bottom=0.7, left=0.8, right=0.8):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_default_font(document, font_name="Aptos", size=10.5):
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            style.font.size = Pt(size)

    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    styles["Title"].font.size = Pt(20)
    styles["Title"].font.bold = True

    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(4)

    styles["Heading 2"].font.size = Pt(11.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(2)

    styles["Heading 3"].font.size = Pt(10.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].paragraph_format.space_before = Pt(6)
    styles["Heading 3"].paragraph_format.space_after = Pt(1)

def add_hyperlink(paragraph, text, url):
    # External hyperlink in python-docx
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_page_break(document):
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")

def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")

def add_callout(document, title, lines):
    table = document.add_table(rows=1 + len(lines), cols=1)
    table.style = "Table Grid"
    hdr = table.cell(0, 0)
    hdr.text = title
    for p in hdr.paragraphs:
        for r in p.runs:
            r.bold = True
    for i, line in enumerate(lines, start=1):
        table.cell(i, 0).text = line

def add_simple_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    document.add_paragraph("")

def add_title_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Native Delivery Handbook")
    r.bold = True
    r.font.size = Pt(20)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Plain-language operating system for client-project organizations")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Designed for 50–70 developers and designers using minimal mandatory tools")

    document.add_paragraph("")
    add_callout(document, "Core recommendation", [
        "Mandatory tools: Notion + GitHub",
        "Optional tools: Cursor, Codex, Perplexity, Antigravity, or other approved AI tools",
        "Optional later: Jira, only if internal coordination truly requires it",
    ])

    document.add_paragraph("")
    p = document.add_paragraph()
    p.add_run("Operating rule: ").bold = True
    p.add_run("If it affects scope, it starts in Notion. If it affects implementation, it must be linked in GitHub.")

    document.add_paragraph("")
    p = document.add_paragraph()
    p.add_run("Official references used in this handbook:").bold = True

    refs = [
        ("GitHub Projects", "https://docs.github.com/en/issues/planning-and-tracking-with-projects/learning-about-projects/about-projects"),
        ("GitHub Issues", "https://docs.github.com/en/issues/tracking-your-work-with-issues/learning-about-issues/quickstart"),
        ("GitHub CODEOWNERS", "https://docs.github.com/en/repositories/managing-your-repositorys-settings-and-features/customizing-your-repository/about-code-owners"),
        ("GitHub branch protection", "https://docs.github.com/en/repositories/configuring-branches-and-merges-in-your-repository/managing-protected-branches"),
        ("Notion help center", "https://www.notion.so/help"),
        ("Notion databases", "https://www.notion.so/help/intro-to-databases"),
        ("Notion API / integrations", "https://developers.notion.com/docs/create-a-notion-integration"),
        ("Jira Software", "https://support.atlassian.com/jira-software-cloud/"),
        ("Jira automation", "https://support.atlassian.com/cloud-automation/docs/jira-automation/"),
        ("Anthropic prompt engineering", "https://docs.anthropic.com/en/docs/build-with-claude/prompt-engineering/overview"),
    ]
    for label, url in refs:
        p = document.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)

def page_1_how_we_build_here(document):
    document.add_heading("Page 1 — How we build here", level=1)

    document.add_paragraph(
        "This organization builds custom client projects. To move faster without creating chaos, "
        "we use a very small operating system. We do not force everyone onto one AI tool. "
        "We standardize the workflow instead."
    )

    document.add_heading("Mandatory tools", level=2)
    add_bullets(document, [
        "Notion = source of truth for requirements, PRDs, approvals, feedback, decisions, and change history",
        "GitHub = source of execution for issues, branches, pull requests, reviews, and releases",
    ])

    document.add_heading("Optional tools", level=2)
    add_bullets(document, [
        "Cursor, Codex, Perplexity, Antigravity, Claude, ChatGPT, or other approved AI tools",
        "Teams may choose their preferred AI tool, but the repo standards remain the same",
    ])

    document.add_heading("Non-negotiable repo standards", level=2)
    add_bullets(document, [
        "Every repo must include AGENTS.md",
        "Every meaningful feature must include PLAN.md",
        "Every PR must link to its issue and the source requirement or PRD",
        "No direct pushes to main",
        "At least one review before merge",
        "Tests are required for new features and important fixes",
    ])

    document.add_heading("Why this works", level=2)
    add_bullets(document, [
        "Low training overhead",
        "Easy to audit",
        "Clear trace from client request to shipped code",
        "AI can be used without losing control",
    ])

    document.add_heading("Learn more", level=2)
    links = [
        ("Notion help center", "https://www.notion.so/help"),
        ("GitHub Issues quickstart", "https://docs.github.com/en/issues/tracking-your-work-with-issues/learning-about-issues/quickstart"),
        ("GitHub Projects", "https://docs.github.com/en/issues/planning-and-tracking-with-projects/learning-about-projects/about-projects"),
    ]
    for label, url in links:
        p = document.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)

def page_2_delivery_workflow(document):
    add_page_break(document)
    document.add_heading("Page 2 — Delivery workflow", level=1)

    document.add_paragraph("Use this workflow for every client request, feature, or change.")

    document.add_heading("Workflow", level=2)
    add_numbered(document, [
        "Client requirement arrives",
        "Create or update a Requirement item in Notion",
        "Review the requirement internally",
        "Create or update the PRD in Notion",
        "Record a change log entry if anything changed",
        "After approval, create a linked GitHub issue or epic",
        "Developer builds using AGENTS.md, PLAN.md, issue, and PRD context",
        "Raise PR with linked issue, linked source, tests, and review notes",
        "Demo the feature",
        "Log new feedback in Notion and create follow-up implementation items as needed",
        "Release and update the source record with the shipped version",
    ])

    document.add_heading("Traceability rule", level=2)
    add_callout(document, "Every implementation item must answer 3 questions", [
        "What client or internal request started this?",
        "Which PRD or approved change does it map to?",
        "Which PR or release shipped it?",
    ])

    document.add_heading("Definition of done", level=2)
    add_bullets(document, [
        "Requirement or change exists in Notion",
        "Implementation work exists in GitHub",
        "PR links back to source record",
        "Tests are attached or referenced",
        "Demo feedback is logged",
        "Release version is recorded",
    ])

def page_3_ai_coding_guide(document):
    add_page_break(document)
    document.add_heading("Page 3 — AI coding guide", level=1)

    document.add_paragraph(
        "Use any approved AI tool you want. The output must still follow the same repo rules."
    )

    document.add_heading("Always give the model these 4 things", level=2)
    add_bullets(document, [
        "AGENTS.md",
        "PLAN.md",
        "The GitHub issue or task",
        "The linked PRD or requirement summary",
    ])

    document.add_heading("What AI is good for", level=2)
    add_bullets(document, [
        "Scaffolding and boilerplate",
        "Test generation",
        "Refactoring",
        "Summarizing requirements",
        "Generating first drafts of docs, PR descriptions, or checklists",
    ])

    document.add_heading("What humans must still verify", level=2)
    add_bullets(document, [
        "Business logic",
        "Security and auth",
        "Edge cases",
        "Data correctness",
        "Client-specific constraints",
    ])

    document.add_heading("Prompt — create AGENTS.md", level=2)
    document.add_paragraph(
        "Create a repo-level AGENTS.md for this project. Keep it plain and practical. "
        "Include: project overview, tech stack, folder structure, coding rules, testing rules, "
        "security rules, review checklist, definition of done, and explicit forbidden shortcuts. "
        "The rules must require no hardcoded secrets, no undocumented client-specific logic, "
        "tests for all new features, validation for all new APIs, migration planning for DB changes, "
        "and human review of all AI-generated code. Use clear markdown headings and bullet points only."
    )

    document.add_heading("Prompt — create PLAN.md", level=2)
    document.add_paragraph(
        "Create a PLAN.md for this feature. Include: feature summary, goals, non-goals, assumptions, "
        "dependencies, architecture notes, implementation phases, testing plan, security checks, risks, "
        "acceptance criteria, and rollback notes. Keep phases small enough that they can be executed in "
        "separate AI chats without losing context."
    )

    document.add_heading("Prompt — build from PLAN.md", level=2)
    document.add_paragraph(
        "Using AGENTS.md and PLAN.md, implement only Phase 1 of the plan. Explain what files will change, "
        "what assumptions you are making, what tests you will add, and what risks you see. Do not implement "
        "later phases. Keep the changes minimal and production-safe."
    )

    document.add_heading("Prompt — self-review", level=2)
    document.add_paragraph(
        "Review this diff against AGENTS.md and PLAN.md. Check for broken business logic, security gaps, "
        "missing tests, hardcoded values, missing validation, unhandled edge cases, migration risks, and "
        "maintainability concerns. Return findings in priority order with exact fixes."
    )

    document.add_heading("Learn more", level=2)
    p = document.add_paragraph(style="List Bullet")
    add_hyperlink(p, "Anthropic prompt engineering guide", "https://docs.anthropic.com/en/docs/build-with-claude/prompt-engineering/overview")

def page_4_templates(document):
    add_page_break(document)
    document.add_heading("Page 4 — Templates", level=1)

    document.add_heading("Required Notion databases", level=2)
    add_simple_table(document,
        ["Database", "Purpose", "Minimum fields"],
        [
            ["Requirements Register", "Capture incoming client asks", "ID, title, client, source, date, status, linked PRD, priority"],
            ["PRD Register", "Store approved product requirements", "PRD ID, feature, version, status, owner, linked requirements"],
            ["Change Log", "Track scope and feedback changes", "Change ID, date, trigger, old scope, new scope, impact, approval"],
            ["Decision Log", "Track key product and engineering decisions", "Decision ID, date, topic, options, final decision, reason"],
        ]
    )

    document.add_heading("PRD page structure", level=2)
    add_bullets(document, [
        "Header: project, feature, owner, status, version",
        "Problem statement",
        "Goals and non-goals",
        "User stories or flows",
        "Functional requirements",
        "Acceptance criteria",
        "Risks and dependencies",
        "Linked change log",
        "Linked implementation items",
        "Links to previous approved versions",
    ])

    document.add_heading("Issue template fields", level=2)
    add_bullets(document, [
        "Title",
        "Problem to solve",
        "Source requirement or PRD link",
        "Acceptance criteria",
        "Risks",
        "Dependencies",
        "Owner",
    ])

    document.add_heading("PR template fields", level=2)
    add_bullets(document, [
        "What changed",
        "Why it changed",
        "Linked issue",
        "Linked PRD or change request",
        "Tests added or run",
        "Security considerations",
        "Rollback notes",
        "Screenshots or demo notes if applicable",
    ])

    document.add_heading("Prompt — convert requirement into PRD draft", level=2)
    document.add_paragraph(
        "Convert the following client requirement into a plain-language PRD draft. "
        "Include problem statement, goals, non-goals, user flow, functional requirements, acceptance criteria, "
        "open questions, assumptions, risks, and a suggested implementation breakdown. Keep it readable for both "
        "technical and non-technical stakeholders."
    )

def page_5_non_negotiable_rules(document):
    add_page_break(document)
    document.add_heading("Page 5 — Non-negotiable rules", level=1)

    add_bullets(document, [
        "No direct pushes to main",
        "Every PR must have at least one review",
        "Every feature must link back to a requirement or PRD",
        "Every important change must create or update a change-log entry",
        "Tests are required for new features and important fixes",
        "No secrets in code or prompts",
        "No undocumented hardcoded client logic",
        "All AI-generated code must be human-reviewed before merge",
        "New APIs require validation and auth review",
        "DB changes require migration planning",
    ])

    document.add_heading("Repo checklist", level=2)
    add_bullets(document, [
        "AGENTS.md present",
        "Issue templates present",
        "PR template present",
        "CODEOWNERS present",
        "Protected main branch enabled",
        "Required checks enabled",
    ])

    document.add_heading("Learn more", level=2)
    links = [
        ("CODEOWNERS", "https://docs.github.com/en/repositories/managing-your-repositorys-settings-and-features/customizing-your-repository/about-code-owners"),
        ("Protected branches", "https://docs.github.com/en/repositories/configuring-branches-and-merges-in-your-repository/managing-protected-branches"),
    ]
    for label, url in links:
        p = document.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)

def page_6_jira_option(document):
    add_page_break(document)
    document.add_heading("Page 6 — Optional Jira version", level=1)

    document.add_paragraph(
        "Use this page only if leadership wants Jira for internal delivery coordination. "
        "Client-facing source of truth should still remain in Notion."
    )

    document.add_heading("Recommended split with Jira", level=2)
    add_bullets(document, [
        "Notion = client-facing requirements, PRDs, approvals, change history",
        "Jira = internal delivery planning, sprint management, team coordination",
        "GitHub = source control, pull requests, reviews, and release links",
    ])

    document.add_heading("Pipeline with Jira", level=2)
    add_numbered(document, [
        "Client request arrives and is logged in Notion",
        "PRD is created or updated in Notion",
        "Approved work is converted into Jira epic/story/task",
        "Jira item links back to Notion source",
        "GitHub branch and PR link back to Jira item",
        "Demo feedback updates Notion change log",
        "Approved follow-up work creates new Jira items",
    ])

    document.add_heading("When Jira is worth it", level=2)
    add_bullets(document, [
        "You need sprint planning across multiple teams",
        "You need delivery reporting for leadership",
        "You need non-engineering stakeholders to manage internal workflow in one place",
    ])

    document.add_heading("Learn more", level=2)
    links = [
        ("Jira Software help", "https://support.atlassian.com/jira-software-cloud/"),
        ("Jira automation", "https://support.atlassian.com/cloud-automation/docs/jira-automation/"),
    ]
    for label, url in links:
        p = document.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)

def page_7_automation(document):
    add_page_break(document)
    document.add_heading("Page 7 — Automation to reduce PRD and change-log work", level=1)

    document.add_paragraph(
        "Do not make people manually rewrite everything. Automate draft generation, but keep human approval before final updates."
    )

    document.add_heading("Recommended automation pattern", level=2)
    add_numbered(document, [
        "Meeting notes, email summary, or client feedback arrives",
        "Automation reads the related Notion page and latest PRD version",
        "AI drafts: PRD delta, change-log entry, decision draft, and suggested implementation items",
        "A human reviews and approves the draft",
        "Approved content is written back into Notion",
        "Optional: linked GitHub issue or Jira item is created automatically",
    ])

    document.add_heading("What should be automated", level=2)
    add_bullets(document, [
        "Requirement summarization",
        "PRD draft generation",
        "Change-log first drafts",
        "Demo-feedback summaries",
        "Decision-log first drafts",
        "Issue creation from approved changes",
    ])

    document.add_heading("What should not be fully automated", level=2)
    add_bullets(document, [
        "Final approval",
        "Scope change acceptance",
        "Security-sensitive implementation decisions",
        "Client commitments",
    ])

    document.add_heading("Prompt — turn meeting notes into change log", level=2)
    document.add_paragraph(
        "Read these meeting notes and compare them against the current PRD summary. "
        "Draft a change-log entry with: date, trigger, old scope, new scope, reason for change, impact on timeline, "
        "impact on risk, and recommended follow-up implementation items. Keep it concise and approval-ready."
    )

    document.add_heading("Learn more", level=2)
    p = document.add_paragraph(style="List Bullet")
    add_hyperlink(p, "Notion integration guide", "https://developers.notion.com/docs/create-a-notion-integration")


# ---------- Build document ----------

doc = Document()
set_default_font(doc, font_name="Aptos", size=10.5)

for section in doc.sections:
    set_page_margins(section)

add_title_page(doc)
page_1_how_we_build_here(doc)
page_2_delivery_workflow(doc)
page_3_ai_coding_guide(doc)
page_4_templates(doc)
page_5_non_negotiable_rules(doc)
page_6_jira_option(doc)
page_7_automation(doc)

doc.save(OUTPUT_FILE)
print(f"Created: {OUTPUT_FILE}")